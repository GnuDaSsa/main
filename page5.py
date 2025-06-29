def run():
    import streamlit as st
    import re
    import os
    from dotenv import load_dotenv
    from openai import OpenAI
    from datetime import datetime
    from docx import Document
    from odf.text import P, Span, H, ListItem
    from odf.opendocument import load
    import io

    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    client = OpenAI(api_key=api_key)

    # 카운터 파일 기반(혹은 DB 연동 유지 시 아래만 남김)
    def 카운터_초기화():
        pass  # 필요시 파일/DB 초기화 코드

    def 카운터_증가():
        if 'press_count' not in st.session_state:
            st.session_state['press_count'] = 1
        else:
            st.session_state['press_count'] += 1
        return st.session_state['press_count']

    # 날짜/카운트 표시용
    now = datetime.now()
    days_kor = ['월', '화', '수', '목', '금', '토', '일']
    weekday_kor = days_kor[now.weekday()]
    formatted_date = now.strftime('%Y.%m.%d') + f'.({weekday_kor})'

    st.title("성남시 생성형 AI 보도자료 작성 서비스(ChatGPT-4o)")

    col1, col2 = st.columns([8.5,1.5]) 
    with col1:
        st.warning("**1. 단어와 간단한 문장을 입력하는 것만으로도 여러분의 아이디어와 정보를 효과적으로 전달할 수 있습니다.**\n\n" 
                   "**2. 복잡한 과정이 필요 없이, 직관적인 사용 방식으로 신속하게 보도자료를 완성해 보세요!**\n\n")

    with col2: 
        count = st.session_state.get('press_count', 0)
        st.error(f"**✔️오늘 {count}회**")

    col3, col4 = st.columns([5.5,4.5])

    # 보도자료 생성 입력 폼
    with col4:
        부서명col = st.columns([1])[0]
        담당부서 = st.text_input("**1.담당부서**", placeholder='4차산업추진국 AI반도체과')
        col5, col6, col7 = st.columns(3)
        with col5:        
            소감 = st.text_input("**2.소감주체**", placeholder='성남시장 OOO')
        with col6:        
            담당자 = st.text_input("**3.주무관**", placeholder='홍길동')
        with col7:            
            연락처 = st.text_input("**4.연락처**", placeholder='031-729-0000')
        내용 = st.text_area("**5.보도자료 핵심반영 내용**", height=170,placeholder='(예시)물놀이장 이용은 초등학생 이하로 성남시민은 신분증 확인후 우선 입장하고, 운영시간은 오전 10시부터 오후 5시까지며 매주 월요일은 휴무입니다. 다양한 물놀이 시설과 안전관리자 간호사 배치')

        # 생성 결과 저장용 변수
        if 'last_press' not in st.session_state:
            st.session_state.last_press = None
        if 'last_title' not in st.session_state:
            st.session_state.last_title = None
        if 'last_download_ready' not in st.session_state:
            st.session_state.last_download_ready = False
        if 'last_titles' not in st.session_state:
            st.session_state.last_titles = []

        col17, col18 = st.columns([2.5,7.5])    
        with col17:
            if st.button("보도자료 생성", use_container_width=True):
                if 담당부서 and 내용 and 담당자 and 연락처:
                    st.session_state['press_ready'] = False
                    with col3:
                        output_placeholder = st.empty()
                        full_text = ""
                        completion = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {"role": "system", "content": f"당신은 배태랑 기자입니다. 입력된 내용을 바탕으로 보도자료 제목 5개를 번호를 붙혀서 추천하고, 이를 '보도자료 추천 제목'이라는 진한글씨로 제목 아래 각각 ''로 묶어서 화면에 보여주세요. 각 추천 제목은 한 줄씩 출력해야합니다. 그 후 추천 제목 중 첫 번째를 주제로 보도자료를 작성하세요. 보도자료 내용은 '{내용}'을 기반으로 주제를 정하여 2500자 이상 풍부하게 작성합니다. 보도자료의 마지막 단락에는 '{소감}의 내용을 출력하고 소감을 풍부하게 작성하고, 담당자와 연락처는 각각 '{담당자}','{연락처}'로 작성합니다. 각 단락은 간결하고 공식적으로 작성해야 하며, 전형적인 보도자료 구조로 문장의 종결은 '이다','된다','했다'로 마무리합니다. "},
                            ],
                            stream=True
                        )
                        for chunk in completion:
                            if hasattr(chunk.choices[0].delta, 'content'):
                                content = chunk.choices[0].delta.content
                                if content:
                                    full_text += content
                                    output_placeholder.write(full_text)
                        output_placeholder.empty()  # 타이핑 끝나면 비우기
                        import re
                        # 1. 추천 제목 5개만 추출
                        title_section = re.search(r"보도자료 추천 제목.*?((?:'[^']+'\s*){5})", full_text, re.DOTALL)
                        if title_section:
                            titles = re.findall(r"'([^']+)'", title_section.group(1))
                        else:
                            titles = []
                        # 2. 본문(실제 보도자료 내용)만 추출
                        body_section = re.split(r"보도자료 추천 제목.*?(?:'[^']+'\s*){5}", full_text, maxsplit=1, flags=re.DOTALL)
                        if len(body_section) > 1:
                            body = body_section[1].strip()
                        else:
                            body = full_text  # fallback
                        st.session_state['press_titles'] = titles
                        st.session_state['press_main_title'] = titles[0] if titles else ''
                        st.session_state['press_full'] = body
                        st.session_state['press_ready'] = True
                        카운터_증가()
                        st.success("보도자료가 성공적으로 생성되었습니다!")
                else:
                    st.error("모든 필드를 입력해주세요.")

        with col18:
            if st.session_state.get('last_download_ready') and st.session_state.last_press:
                title = st.session_state.last_title.replace(" ", "_")[:15]
                Press_path = "SeongNam_Press.odt"  # 수정할 파일 경로
                try:
                    with open(Press_path, "rb") as file:
                        odt_file = file.read()
                    doc = load(io.BytesIO(odt_file))
                    def replace_text_in_node(node, search_text, replace_text):
                        if node.nodeType == node.TEXT_NODE:
                            if search_text in node.data:
                                node.data = node.data.replace(search_text, str(replace_text))
                        else:
                            for child in node.childNodes:
                                replace_text_in_node(child, search_text, replace_text)
                    def replace_text_in_elements(doc, search_text, replace_text):
                        elements = [P, Span, H, ListItem]
                        for elem in elements:
                            for node in doc.getElementsByType(elem):
                                replace_text_in_node(node, search_text, replace_text)
                    def add_paragraph(doc, text):
                        paragraph = P(text=text)
                        doc.text.addElement(paragraph)
                    replace_text_in_elements(doc, "입력부서", 담당부서)
                    replace_text_in_elements(doc, "입력담당자", 담당자)
                    replace_text_in_elements(doc, "입력연락처", 연락처)
                    replace_text_in_elements(doc, "입력제목", st.session_state.last_title)
                    replace_text_in_elements(doc, "입력일자", formatted_date)
                    add_paragraph(doc, " ")
                    입력본문 = st.session_state.last_press
                    paragraphs = 입력본문.strip().split('\n')
                    for i, paragraph in enumerate(paragraphs):
                        add_paragraph(doc, paragraph)
                    odt_out_file = io.BytesIO()
                    doc.save(odt_out_file)
                    odt_out_file.seek(0)
                    st.download_button(
                        label=f"{title}보도자료 다운로드",
                        data=odt_out_file,
                        file_name=f"{title}.odt",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"문서 파일 생성 중 오류 발생: {e}")
            else:
                st.info("보도자료를 먼저 생성해 주세요.")

    # col3: 추천 제목 5개와 본문을 한 번만 정확히 출력
    with col3:
        if st.session_state.get('press_ready'):
            if st.session_state.get('press_titles'):
                st.markdown("**AI 추천 보도자료 제목 5개:**")
                for i, t in enumerate(st.session_state['press_titles'], 1):
                    st.markdown(f"{i}. {t}")
            st.subheader(st.session_state.get('press_main_title', ''))
            st.write(st.session_state.get('press_full', ''))
        else:
            st.info("보도자료를 생성하면 이곳에 결과가 표시됩니다.")
